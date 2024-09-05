from docx import Document

# Create a new Document
doc = Document()

# Add a Title
doc.add_heading('Sales Data Analysis Tool Documentation', level=1)

# Add the Introduction
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "This document provides an overview of the Sales Data Analysis Tool. The tool is built using Python and includes functionalities for loading, analyzing, and visualizing sales data. The application features a graphical user interface (GUI) created with Tkinter, which integrates various types of visualizations including bar charts, line charts, and pie charts."
)

# Add Code Section
doc.add_heading('Code', level=2)
doc.add_paragraph(
    "Below is the complete code for the Sales Data Analysis Tool:"
)

# Add the code block
code = """
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from tkinter import Tk, filedialog, Button, Label, Frame, Text, Scrollbar, LEFT, RIGHT, Y, END
from tkinter import ttk
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.widgets import Cursor
import matplotlib.dates as mdates

def load_and_clean_data(file_path):
    df = pd.read_excel(file_path)
    df['Order Date'] = pd.to_datetime(df['Order Date'])
    df['Total Revenue'] = pd.to_numeric(df['Total Revenue'], errors='coerce')
    df = df.dropna()
    df['Total Revenue'] = df['Total Revenue'] / 1_000_000
    return df

def analyze_data(df):
    total_sales = df['Total Revenue'].sum()
    avg_sales = df['Total Revenue'].mean()
    sales_by_product = df.groupby('Item Type')['Total Revenue'].sum().reset_index()
    return total_sales, avg_sales, sales_by_product

def generate_visualizations(df, sort_order='desc'):
    sns.set_style("darkgrid", {"axes.facecolor": "#2c3e50"})
    bar_data = df.groupby('Item Type')['Total Revenue'].sum().reset_index()
    bar_data = bar_data.sort_values(by='Total Revenue', ascending=(sort_order == 'asc'))
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = sns.barplot(x='Item Type', y='Total Revenue', data=bar_data, ax=ax, palette="viridis")
    ax.set_title('Total Sales by Product', color="#e0e0e0")
    ax.set_facecolor('#2c3e50')
    ax.tick_params(colors='#e0e0e0')
    def on_hover_bar(event):
        if event.inaxes == ax:
            for bar in bars.patches:
                if bar.contains(event)[0]:
                    height = bar.get_height()
                    ax.annotate(f'{height:,.2f}M',
                                (bar.get_x() + bar.get_width() / 2, height),
                                xytext=(0, 5),
                                textcoords='offset points', 
                                ha='center', va='bottom',
                                color='#e0e0e0')
                    fig.canvas.draw_idle()
                    break
            else:
                ax.annotate("", xy=(0,0), xytext=(0,0))
                fig.canvas.draw_idle()
    fig.canvas.mpl_connect("motion_notify_event", on_hover_bar)
    line_fig, line_ax = plt.subplots(figsize=(10, 6))
    line_data = df.set_index('Order Date').resample('Y').sum()['Total Revenue']
    line_ax.plot(line_data.index, line_data.values, color='#ff6f61')
    line_ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y'))
    line_ax.xaxis.set_major_locator(mdates.YearLocator())
    line_ax.set_title('Sales Over Time', color="#e0e0e0")
    line_ax.set_ylabel('Total Revenue (Millions)', color="#e0e0e0")
    line_ax.set_xlabel('Year', color="#e0e0e0")
    line_ax.set_facecolor('#2c3e50')
    line_ax.tick_params(colors='#e0e0e0')
    cursor = Cursor(line_ax, useblit=True, color='red', linewidth=1)
    def on_hover_line(event):
        if event.inaxes == line_ax:
            xdata = event.xdata
            ydata = event.ydata
            line_ax.annotate(f'{ydata:,.2f}M',
                             (xdata, ydata),
                             xytext=(0, 5),
                             textcoords='offset points',
                             ha='center', va='bottom',
                             color='#e0e0e0')
            line_fig.canvas.draw_idle()
    line_fig.canvas.mpl_connect("motion_notify_event", on_hover_line)
    pie_fig, pie_ax = plt.subplots(figsize=(7, 7))
    sales_by_product = df.groupby('Item Type')['Total Revenue'].sum()
    wedges, texts, autotexts = pie_ax.pie(sales_by_product, labels=sales_by_product.index, autopct='%1.1f%%',
    startangle=140, colors=sns.color_palette("magma"),
    textprops={'color': '#e0e0e0'})
    pie_ax.set_title('Sales Distribution by Product', color="#e0e0e0")
    pie_ax.set_facecolor('#2c3e50')
    def on_hover_pie(event):
        if event.inaxes == pie_ax:
            for wedge in wedges:
                if wedge.contains(event)[0]:
                    wedge.set_edgecolor('yellow')
                    wedge.set_linewidth(2)
                    pie_fig.canvas.draw_idle()
                    break
            else:
                for wedge in wedges:
                    wedge.set_edgecolor('black')
                    wedge.set_linewidth(1)
                pie_fig.canvas.draw_idle()
    pie_fig.canvas.mpl_connect("motion_notify_event", on_hover_pie)
    plt.tight_layout()
    fig.patch.set_facecolor('#2c3e50')
    line_fig.patch.set_facecolor('#2c3e50')
    pie_fig.patch.set_facecolor('#2c3e50')
    return fig, line_fig, pie_fig

def generate_summary_report(total_sales, avg_sales, sales_by_product):
    report = f"Summary Report\\n"
    report += f"====================\\n"
    report += f"Total Sales: ${total_sales:,.2f} Million\\n"
    report += f"Average Sales: ${avg_sales:,.2f} Million\\n\\n"
    report += "Sales by Product:\\n"
    report += sales_by_product.to_markdown(index=False)
    return report

def show_gui():
    def on_load():
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            df = load_and_clean_data(file_path)
            total_sales, avg_sales, sales_by_product = analyze_data(df)
            report = generate_summary_report(total_sales, avg_sales, sales_by_product)
            text_box.delete(1.0, END)
            text_box.insert(END, report)
            fig_dict = {
                'bar_desc': generate_visualizations(df, 'desc')[0],
                'bar_asc': generate_visualizations(df, 'asc')[0],
                'line': generate_visualizations(df, 'desc')[1],
                'pie': generate_visualizations(df, 'desc')[2]
            }
            def update_canvas(fig):
                for widget in canvas_frame.winfo_children():
                    widget.destroy()
                canvas = FigureCanvasTkAgg(fig, master=canvas_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(side=LEFT, fill='both', expand=True)
            def show_bar_chart_asc():
                update_canvas(fig_dict['bar_asc'])
            def show_bar_chart_desc():
                update_canvas(fig_dict['bar_desc'])
            def show_line_chart():
                update_canvas(fig_dict['line'])
            def show_pie_chart():
                update_canvas(fig_dict['pie'])
            btn_bar_asc.config(command=show_bar_chart_asc)
            btn_bar_desc.config(command=show_bar_chart_desc)
            btn_line.config(command=show_line_chart)
            btn_pie.config(command=show_pie_chart)
    root = Tk()
    root.title("Sales Data Analysis")
    root.geometry("1000x700")
    root.configure(bg='lightskyblue1')
    title_label = Label(root, text="Sales Data Analysis", font=("Arial", 20), bg='lightskyblue1')
    title_label.pack(pady=10)
    btn_frame = Frame(root, bg='lightskyblue1')
    btn_frame.pack(pady=10)
    btn_load = Button(btn_frame, text="Load Data", command=on_load, font=("Arial", 14))
    btn_load.pack(side=LEFT, padx=10)
    btn_bar_desc = Button(btn_frame, text="Bar Chart (Desc)", font=("Arial", 14))
    btn_bar_desc.pack(side=LEFT, padx=10)
    btn_bar_asc = Button(btn_frame, text="Bar Chart (Asc)", font=("Arial", 14))
    btn_bar_asc.pack(side=LEFT, padx=10)
    btn_line = Button(btn_frame, text="Line Chart", font=("Arial", 14))
    btn_line.pack(side=LEFT, padx=10)
    btn_pie = Button(btn_frame, text="Pie Chart", font=("Arial", 14))
    btn_pie.pack(side=LEFT, padx=10)
    text_frame = Frame(root, bg='lightskyblue1')
    text_frame.pack(fill=BOTH, expand=True)
    text_box = Text(text_frame, wrap=WORD, font=("Arial", 12))
    text_box.pack(side=LEFT, fill=BOTH, expand=True)
    scrollbar = Scrollbar(text_frame, command=text_box.yview)
    scrollbar.pack(side=RIGHT, fill=Y)
    text_box.config(yscrollcommand=scrollbar.set)
    canvas_frame = Frame(root, bg='lightskyblue1')
    canvas_frame.pack(fill=BOTH, expand=True)
    root.mainloop()

if __name__ == "__main__":
    show_gui()
"""
doc.add_paragraph(code, style='Normal')

# Save the Document
doc.save('Sales_Data_Analysis_Documentation.docx')
